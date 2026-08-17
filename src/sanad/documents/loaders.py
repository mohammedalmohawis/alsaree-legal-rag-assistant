"""File parsing: bytes in, structured :class:`TextBlock` list out.

Both loaders preserve *structure*, not just characters, because everything
downstream depends on it — the chunker needs paragraph and heading boundaries,
and citations need the page a block came from.

PDF pages are 1-based, matching what a reader sees in a viewer. DOCX has no
page concept before rendering, so those blocks carry ``page=None`` and the
resulting citations say so explicitly instead of guessing.
"""

from __future__ import annotations

import io
import os
from collections.abc import Iterator, Sequence
from typing import Callable

from sanad.documents.models import ParsedDocument, SourceDocument, TextBlock
from sanad.documents.structure import detect_label, is_heading
from sanad.errors import DocumentReadError, EmptyDocumentError, UnsupportedFormatError
from sanad.utils.text import collapse_whitespace

try:  # PyMuPDF renamed its import from ``fitz`` to ``pymupdf`` in v1.24.
    import pymupdf
except ImportError:  # pragma: no cover - exercised only on older PyMuPDF
    import fitz as pymupdf  # type: ignore[no-redef]

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SUPPORTED_EXTENSIONS = (".pdf", ".docx")


def _make_blocks(
    lines: Sequence[str],
    page: int | None,
    style_hint: str | None = None,
) -> list[TextBlock]:
    """Turn raw lines into heading and paragraph blocks.

    Consecutive body lines are merged into one paragraph block so that a
    sentence wrapped across three PDF lines stays a single unit; heading lines
    always stand alone so the chunker can start a new section at them.
    """
    blocks: list[TextBlock] = []
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        text = collapse_whitespace(" ".join(buffer))
        if text:
            blocks.append(TextBlock(text=text, page=page, label=detect_label(text)))
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        if is_heading(stripped, style_hint=style_hint):
            flush()
            blocks.append(
                TextBlock(
                    text=collapse_whitespace(stripped),
                    page=page,
                    is_heading=True,
                    label=detect_label(stripped),
                    kind="heading",
                )
            )
        else:
            buffer.append(stripped)

    flush()
    return blocks


def load_pdf(filename: str, data: bytes) -> tuple[list[TextBlock], int]:
    """Extract page-attributed blocks from a PDF.

    Returns ``(blocks, page_count)``. PyMuPDF's ``"blocks"`` mode is used rather
    than plain text because it returns layout blocks in reading order, which
    preserves the paragraph boundaries the chunker relies on.
    """
    blocks: list[TextBlock] = []
    try:
        with pymupdf.open(stream=data, filetype="pdf") as pdf:
            page_count = pdf.page_count
            for page_number, page in enumerate(pdf, start=1):
                raw_blocks = page.get_text("blocks") or []
                # Sort top-to-bottom, then left-to-right, to recover reading order.
                for raw in sorted(raw_blocks, key=lambda b: (round(b[1], 1), round(b[0], 1))):
                    # Tuple layout: (x0, y0, x1, y1, text, block_no, block_type).
                    if len(raw) >= 7 and raw[6] != 0:
                        continue  # image block, no text to extract
                    text = raw[4] if len(raw) > 4 else ""
                    if not isinstance(text, str) or not text.strip():
                        continue
                    blocks.extend(_make_blocks(text.splitlines(), page_number))
    except DocumentReadError:
        raise
    except Exception as error:  # PyMuPDF raises a wide range of exception types.
        raise DocumentReadError(filename, str(error)) from error
    return blocks, page_count


def _iter_docx_body(document: DocxDocument) -> Iterator[object]:
    """Yield paragraphs and tables in true document order.

    ``document.paragraphs`` and ``document.tables`` are separate sequences, so
    reading them one after the other would move every table to the end of the
    document. Walking the body XML keeps a payment schedule next to the clause
    that introduces it.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_to_lines(table: Table) -> list[str]:
    """Render a table as one pipe-delimited line per row.

    Keeping a row on a single line preserves the association between a label and
    its value ("Monthly fee | SAR 12,000"), which matters more for retrieval
    than reproducing the visual grid.
    """
    lines: list[str] = []
    for row in table.rows:
        cells = [collapse_whitespace(cell.text) for cell in row.cells]
        # Word merges cells by repeating them; collapse the duplicates.
        deduped: list[str] = []
        for cell in cells:
            if not deduped or deduped[-1] != cell:
                deduped.append(cell)
        rendered = " | ".join(cell for cell in deduped if cell)
        if rendered:
            lines.append(rendered)
    return lines


def load_docx(filename: str, data: bytes) -> tuple[list[TextBlock], None]:
    """Extract blocks from a DOCX, preserving heading styles and tables.

    The page count is ``None`` because pagination does not exist until Word
    renders the file; downstream code treats that as "cite by section only".
    """
    blocks: list[TextBlock] = []
    try:
        document = DocxDocument(io.BytesIO(data))
        for item in _iter_docx_body(document):
            if isinstance(item, Table):
                for line in _table_to_lines(item):
                    blocks.append(
                        TextBlock(
                            text=line,
                            page=None,
                            label=detect_label(line),
                            kind="table",
                        )
                    )
                continue

            paragraph = item  # type: ignore[assignment]
            text = collapse_whitespace(paragraph.text)
            if not text:
                continue
            style_name = getattr(getattr(paragraph, "style", None), "name", None)
            blocks.extend(_make_blocks([text], page=None, style_hint=style_name))
    except DocumentReadError:
        raise
    except Exception as error:  # python-docx raises package-specific errors.
        raise DocumentReadError(filename, str(error)) from error
    return blocks, None


#: Extension -> loader. Adding a format means adding one entry here.
LOADERS: dict = {
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(doc_id: str, filename: str, data: bytes) -> ParsedDocument:
    """Parse one uploaded file into a :class:`ParsedDocument` without chunks.

    Raises :class:`UnsupportedFormatError` for unknown extensions and
    :class:`EmptyDocumentError` when a file parses but yields no text — the
    signature of a scanned PDF with no OCR layer.
    """
    extension = os.path.splitext(filename)[1].lower()
    loader: Callable[[str, bytes], tuple[list[TextBlock], int | None]] | None
    loader = LOADERS.get(extension)
    if loader is None:
        raise UnsupportedFormatError(filename, extension or "unknown")

    blocks, page_count = loader(filename, data)
    blocks = [block for block in blocks if block.text.strip()]
    if not blocks:
        raise EmptyDocumentError(filename)

    char_count = sum(len(block.text) for block in blocks)
    document = SourceDocument(
        doc_id=doc_id,
        filename=filename,
        file_type=extension.lstrip("."),
        char_count=char_count,
        page_count=page_count,
        block_count=len(blocks),
    )
    return ParsedDocument(document=document, blocks=blocks)


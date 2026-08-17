"""Shared fixtures.

The whole suite runs offline. Two fakes make that possible:

``FakeTransport``
    A scriptable stand-in for the HTTP layer, so the real
    :class:`~sanad.llm.gemini.GeminiClient` — retries, batching, response
    parsing, key redaction — is exercised without a network or an API key.

``bag_of_words``
    A deterministic embedding over a fixed vocabulary. Because it is a real
    (if crude) semantic representation, retrieval assertions test genuine
    ranking behaviour rather than the shape of random numbers.

Document fixtures are generated at run time with PyMuPDF and python-docx, so no
binary files are committed to the repository.
"""

from __future__ import annotations

import io
import math
from collections.abc import Mapping, Sequence
from typing import Any

import pytest

from sanad.config import AppSettings, ChunkingSettings, GeminiSettings
from sanad.llm.transport import NetworkError, TransportResponse

# --- deterministic embeddings -------------------------------------------------

#: Fixed vocabulary. Terms are drawn from the sample contracts below so that
#: related passages genuinely land near each other in the vector space.
VOCABULARY = [
    "termination", "terminate", "notice", "days", "payment", "fee", "salary",
    "party", "parties", "employer", "employee", "contract", "agreement",
    "governing", "law", "jurisdiction", "confidential", "penalty", "renewal",
    "sar", "riyadh", "annual", "leave", "probation", "deliver", "goods",
]


def bag_of_words(text: str) -> list[float]:
    """A unit-length count vector over :data:`VOCABULARY`."""
    lowered = text.lower()
    values = [float(lowered.count(term)) for term in VOCABULARY]
    magnitude = math.sqrt(sum(value * value for value in values))
    if magnitude == 0:
        # Keep every vector unit-length so cosine similarity stays well defined.
        return [1.0 / math.sqrt(len(VOCABULARY))] * len(VOCABULARY)
    return [value / magnitude for value in values]


# --- HTTP fake ----------------------------------------------------------------


class FakeTransport:
    """Records requests and replies from a script or an auto-responder."""

    def __init__(
        self,
        responses: Sequence[Any] | None = None,
        *,
        text_reply: str = "Answer grounded in the documents. [S1]",
    ) -> None:
        self.requests: list[dict[str, Any]] = []
        self._responses = list(responses or [])
        self._text_reply = text_reply

    def post(
        self,
        url: str,
        *,
        headers: Mapping[str, str],
        json: Mapping[str, Any],
        timeout: int,
    ) -> TransportResponse:
        self.requests.append(
            {"url": url, "headers": dict(headers), "json": dict(json), "timeout": timeout}
        )

        if self._responses:
            nxt = self._responses.pop(0)
            if isinstance(nxt, Exception):
                raise nxt
            return nxt

        return self._auto(url, json)

    def _auto(self, url: str, payload: Mapping[str, Any]) -> TransportResponse:
        """Answer embedding and generation calls with plausible payloads."""
        if url.endswith(":batchEmbedContents"):
            texts = [
                request["content"]["parts"][0]["text"] for request in payload["requests"]
            ]
            return TransportResponse(
                200, {"embeddings": [{"values": bag_of_words(text)} for text in texts]}
            )
        if url.endswith(":embedContent"):
            text = payload["content"]["parts"][0]["text"]
            return TransportResponse(200, {"embedding": {"values": bag_of_words(text)}})
        return TransportResponse(
            200,
            {"candidates": [{"content": {"parts": [{"text": self._text_reply}]}}]},
        )

    # --- assertions helpers ---------------------------------------------------

    @property
    def call_count(self) -> int:
        return len(self.requests)

    def calls_to(self, suffix: str) -> list[dict[str, Any]]:
        return [request for request in self.requests if request["url"].endswith(suffix)]


def ok(payload: dict[str, Any]) -> TransportResponse:
    return TransportResponse(200, payload)


def failure(status: int, message: str = "boom") -> TransportResponse:
    return TransportResponse(status, {"error": {"message": message}})


def text_response(text: str) -> TransportResponse:
    return ok({"candidates": [{"content": {"parts": [{"text": text}]}}]})


# --- settings -----------------------------------------------------------------


@pytest.fixture
def gemini_settings() -> GeminiSettings:
    return GeminiSettings(
        api_key="test-key-do-not-use",
        max_retries=3,
        backoff_seconds=0.0,
        embedding_batch_size=4,
    )


@pytest.fixture
def app_settings(gemini_settings: GeminiSettings) -> AppSettings:
    return AppSettings(
        gemini=gemini_settings,
        chunking=ChunkingSettings(target_chars=400, max_chars=700, min_chars=80, overlap_chars=60),
    )


@pytest.fixture
def transport() -> FakeTransport:
    return FakeTransport()


@pytest.fixture
def no_sleep():
    """A sleeper that records delays instead of spending them."""
    delays: list[float] = []
    return delays, delays.append


# --- document fixtures --------------------------------------------------------

EMPLOYMENT_CLAUSES = [
    "EMPLOYMENT AGREEMENT",
    "This Employment Agreement is made between Alpha Trading Company, the "
    "Employer, and Sara Al-Otaibi, the Employee.",
    "Section 1. Term",
    "The initial term is twenty-four months commencing 1 March 2026, subject to "
    "a probation period of ninety days.",
    "Section 2. Payment",
    "The Employer shall pay a monthly salary of SAR 18,500 on the last working "
    "day of each month.",
    "Section 3. Termination",
    "Either party may terminate this contract by giving sixty days written "
    "notice. The Employer may terminate immediately for gross misconduct.",
    "Section 4. Governing Law",
    "This agreement is governed by the laws of the Kingdom of Saudi Arabia and "
    "the courts of Riyadh have exclusive jurisdiction.",
]


def build_pdf(pages: Sequence[Sequence[str]]) -> bytes:
    """Render text into a real PDF, one list of lines per page."""
    import pymupdf

    document = pymupdf.open()
    for lines in pages:
        page = document.new_page()
        offset = 72
        for line in lines:
            page.insert_text((72, offset), line, fontsize=11)
            offset += 22
    data = document.tobytes()
    document.close()
    return data


def build_docx(paragraphs: Sequence[tuple], tables: Sequence[Sequence[Sequence[str]]] = ()) -> bytes:
    """Build a real DOCX. ``paragraphs`` is a sequence of ``(text, style)`` pairs."""
    from docx import Document

    document = Document()
    for text, style in paragraphs:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    for rows in tables:
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for row_index, row in enumerate(rows):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def employment_pdf() -> bytes:
    return build_pdf([EMPLOYMENT_CLAUSES[:7], EMPLOYMENT_CLAUSES[7:]])


@pytest.fixture
def employment_docx() -> bytes:
    return build_docx(
        [
            ("Employment Agreement", "Title"),
            ("Section 1. Term", "Heading 1"),
            ("The initial term is twenty-four months.", None),
            ("Section 2. Payment", "Heading 1"),
            ("The Employer shall pay a monthly salary of SAR 18,500.", None),
        ],
        tables=[[["Item", "Amount"], ["Monthly salary", "SAR 18,500"]]],
    )


@pytest.fixture
def arabic_docx() -> bytes:
    return build_docx(
        [
            ("عقد عمل", "Title"),
            ("المادة 1 المدة", "Heading 1"),
            ("مدة هذا العقد أربعة وعشرون شهراً تبدأ من تاريخ الالتحاق بالعمل.", None),
            ("المادة 2 الأجر", "Heading 1"),
            ("يستحق الموظف أجراً شهرياً قدره 18,500 ريال سعودي.", None),
        ]
    )


__all__ = [
    "EMPLOYMENT_CLAUSES",
    "FakeTransport",
    "NetworkError",
    "VOCABULARY",
    "bag_of_words",
    "build_docx",
    "build_pdf",
    "failure",
    "ok",
    "text_response",
]
